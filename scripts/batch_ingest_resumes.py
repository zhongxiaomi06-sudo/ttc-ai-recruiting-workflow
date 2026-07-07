#!/usr/bin/env python3
"""Batch-read local resumes and optionally send them to TTC Daemon.

Default mode is dry-run: it scans and extracts text only. Add --send to enqueue
records through /ingest/link, which works for both local and server Daemon.
"""
import argparse
import json
import sys
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

import requests


DEFAULT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
RESUME_HINTS = ("简历", "resume", "cv", "候选人")


def find_resume_files(roots: Iterable[Path], extensions: set[str]) -> List[Path]:
    files: List[Path] = []
    for root in roots:
        root = root.expanduser()
        if root.is_file() and root.suffix.lower() in extensions:
            files.append(root)
            continue
        if not root.exists() or not root.is_dir():
            continue
        for path in root.rglob("*"):
            if not path.is_file() or path.suffix.lower() not in extensions:
                continue
            name = path.name.lower()
            if any(hint.lower() in name for hint in RESUME_HINTS) or "简历数据" in str(path):
                files.append(path)
    return sorted(set(files), key=lambda p: str(p))


def extract_pdf(path: Path) -> str:
    try:
        import pdfplumber

        chunks: List[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append(text)
        if chunks:
            return "\n\n".join(chunks)
    except Exception:
        pass

    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        chunks = [(page.extract_text() or "") for page in reader.pages]
        return "\n\n".join(t for t in chunks if t.strip())
    except Exception as e:
        raise RuntimeError(f"PDF parse failed: {e}") from e


def extract_docx(path: Path) -> str:
    try:
        import docx
    except Exception as e:
        raise RuntimeError("python-docx is required for .docx files") from e

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    raise RuntimeError(f"Unsupported extension: {suffix}")


def post_resume(daemon_url: str, token: str, path: Path, text: str) -> Dict[str, Any]:
    headers = {"Content-Type": "application/json"}
    if token:
        headers["X-TTC-Token"] = token
    payload = {
        "source_type": "candidate_resume",
        "source_url": str(path),
        "title": path.stem,
        "raw_text": text,
        "markdown": text,
    }
    session = requests.Session()
    session.trust_env = False
    resp = session.post(
        daemon_url.rstrip("/") + "/ingest/link",
        headers=headers,
        json=payload,
        timeout=30,
    )
    try:
        body = resp.json()
    except Exception:
        body = {"text": resp.text[:500]}
    return {"status_code": resp.status_code, "body": body}


def main() -> int:
    parser = argparse.ArgumentParser(description="批量读取本机简历并提交到 TTC Daemon")
    parser.add_argument(
        "--root",
        action="append",
        default=[],
        help="扫描目录或文件，可重复。默认扫描 ./简历数据",
    )
    parser.add_argument("--daemon-url", default="http://127.0.0.1:8766")
    parser.add_argument("--token", default="")
    parser.add_argument("--send", action="store_true", help="真的提交到 Daemon；默认只 dry-run")
    parser.add_argument("--limit", type=int, default=0, help="最多处理多少份，0 表示不限")
    parser.add_argument("--min-chars", type=int, default=120, help="低于该字数视为解析失败")
    parser.add_argument("--report", default="", help="输出 JSONL 报告路径")
    parser.add_argument(
        "--extensions",
        default=",".join(sorted(DEFAULT_EXTENSIONS)),
        help="允许扩展名，逗号分隔，例如 .pdf,.docx,.txt",
    )
    args = parser.parse_args()

    roots = [Path(p) for p in args.root] or [Path("简历数据")]
    extensions = {e.strip().lower() for e in args.extensions.split(",") if e.strip()}
    files = find_resume_files(roots, extensions)
    if args.limit:
        files = files[: args.limit]

    report_fh: Optional[Any] = None
    if args.report:
        report_path = Path(args.report)
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_fh = report_path.open("w", encoding="utf-8")

    stats = {"found": len(files), "parsed": 0, "sent": 0, "failed": 0}
    print(f"found={stats['found']} send={args.send} daemon={args.daemon_url}")

    try:
        for path in files:
            item: Dict[str, Any] = {"path": str(path), "ok": False, "sent": False}
            try:
                text = extract_text(path)
                item["chars"] = len(text)
                if len(text.strip()) < args.min_chars:
                    raise RuntimeError(f"too little text extracted: {len(text.strip())} chars")
                stats["parsed"] += 1
                item["ok"] = True
                if args.send:
                    result = post_resume(args.daemon_url, args.token, path, text)
                    item["send_result"] = result
                    if 200 <= result["status_code"] < 300 and result["body"].get("ok"):
                        item["sent"] = True
                        stats["sent"] += 1
                    else:
                        stats["failed"] += 1
                print(f"OK {path} chars={item['chars']} sent={item['sent']}")
            except Exception as e:
                item["error"] = str(e)
                stats["failed"] += 1
                print(f"FAIL {path}: {e}", file=sys.stderr)
            if report_fh:
                report_fh.write(json.dumps(item, ensure_ascii=False) + "\n")
    finally:
        if report_fh:
            report_fh.close()

    print("summary", json.dumps(stats, ensure_ascii=False))
    return 0 if stats["failed"] == 0 or stats["parsed"] > 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
