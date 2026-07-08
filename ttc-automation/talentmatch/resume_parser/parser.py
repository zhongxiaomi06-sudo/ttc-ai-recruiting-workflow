"""Resume parser core - PDF/DOCX/IMG extraction + LLM structured output"""
from __future__ import annotations
import json
import os
import re
from pathlib import Path
from typing import Optional
from loguru import logger

from .models import ResumeOutput
from .llm_utils import get_llm, reset_llm

# ── Text extraction ──────────────────────────────────────

def extract_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        if text.strip():
            return text.strip()
    except Exception as e:
        logger.warning(f"pypdf failed for {path}: {e}")
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            text = "\n".join(p.extract_text() or "" for p in pdf.pages)
            return text.strip()
    except Exception as e:
        logger.warning(f"pdfplumber failed for {path}: {e}")
        return ""


def extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    tables_text.append(row_text)
        return "\n".join(paragraphs + tables_text).strip()
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def extract_image(path: str) -> str:
    for extractor_name, extractor_fn in [
        ("pytesseract", lambda: _ocr_tesseract(path)),
        ("paddleocr", lambda: _ocr_paddle(path)),
    ]:
        try:
            result = extractor_fn()
            if result:
                return result
        except Exception:
            pass
    return ""


def _ocr_tesseract(path):
    import pytesseract
    from PIL import Image
    return pytesseract.image_to_string(Image.open(path), lang='chi_sim+eng').strip()


def _ocr_paddle(path):
    from paddleocr import PaddleOCR
    ocr = PaddleOCR(use_angle_cls=True, lang='ch')
    result = ocr.ocr(path, cls=True)
    texts = [line[1][0] for line in result[0]] if result and result[0] else []
    return "\n".join(texts).strip()


def extract_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    extractors = {
        ".pdf": extract_pdf, ".docx": extract_docx, ".doc": extract_docx,
        ".png": extract_image, ".jpg": extract_image, ".jpeg": extract_image, ".webp": extract_image,
    }
    if ext in (".txt", ".md", ".tex", ".rtf"):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    extractor = extractors.get(ext)
    if extractor:
        return extractor(path)
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read().strip()
            if len(text) > 50:
                return text
    except Exception:
        pass
    return ""


# ── LLM structured extraction ──────────────────────────────

EXTRACTION_PROMPT = """你是资深猎头简历分析师。从简历原文提取结构化信息，严格按JSON格式返回。

规则：
1. skills_classified 必须按领域分类
2. education 中 985/211/QS50 字段需判断
3. work_experience 中 company_tier: FAANG/大厂/独角兽/上市公司/中小型
4. career_stability: 稳定/一般/频繁跳槽
5. tech_depth: 浅/中/深
6. highlights 提取3-5个最亮眼卖点
7. summary 50字以内概括候选人核心价值

返回JSON：
{"candidate_name":"","email":"","phone":"","wechat":"","location":"","current_role":"","current_company":"","years_experience":0,"education":[{"degree":"","institution":"","major":"","start_date":"","end_date":"","is_985":false,"is_211":false,"is_qs50":false}],"work_experience":[{"position":"","company":"","start_date":"","end_date":"","is_current":false,"description":"","company_tier":"","level":"","team_size":null}],"projects":[{"name":"","role":"","description":"","tech_stack":[],"impact":"","scale":""}],"skills":[],"skills_classified":{"ai_ml":[],"backend":[],"frontend":[],"cloud_devops":[],"data":[],"mobile":[],"product":[],"management":[],"other":[]},"certifications":[],"highlights":[],"career_stability":"","tech_depth":"","industry_tags":[],"role_level":"","summary":""}"""


class ResumeParser:
    def __init__(self, api_key: Optional[str] = None, base_url: Optional[str] = None, model: str = ""):
        self.api_key = api_key
        self.base_url = base_url
        self.model = model

    def parse_file(self, file_path: str) -> ResumeOutput:
        logger.info(f"Parsing: {file_path}")
        raw_text = extract_text(file_path)
        if not raw_text:
            return ResumeOutput(source_file=file_path)

        structured_data = self._llm_extract(raw_text)
        ats = self._compute_ats(raw_text, structured_data)
        structured_data["raw_text"] = raw_text[:2000]
        structured_data["source_file"] = file_path
        structured_data["ats_score"] = ats

        try:
            return ResumeOutput(**structured_data)
        except Exception as e:
            logger.error(f"Pydantic validation failed: {e}")
            return ResumeOutput(raw_text=raw_text[:2000], source_file=file_path,
                               summary=f"解析部分失败: {str(e)[:100]}", ats_score=ats)

    def _llm_extract(self, text: str) -> dict:
        """Extract structured data using LLM with rule-based fallback."""
        client, model = get_llm()
        if not client:
            logger.warning("No LLM client available, using rule-based fallback")
            return self._rule_extract(text)
        use_model = self.model or model
        try:
            resp = client.chat.completions.create(
                model=use_model,
                messages=[
                    {"role": "system", "content": EXTRACTION_PROMPT},
                    {"role": "user", "content": f"分析简历原文：\n\n{text[:10000]}"}
                ],
                temperature=0.1, max_tokens=4096,
                response_format={"type": "json_object"}
            )
            return json.loads(resp.choices[0].message.content)
        except Exception as e:
            logger.error(f"LLM extraction failed: {e}")
            reset_llm()
            return {}

    def _compute_ats(self, raw_text: str, structured: dict) -> float:
        """猎头视角 ATS 评分 — 更接近真实人才评估的分布。
        
        满分 100，维度权重:
          - 教育背景 (15): 学历层次 + 学校档次 + 专业匹配度
          - 经验匹配 (25): 年资适当 + 公司档次 + 稳定性 + 职级
          - 技能深度 (25): AI/技术硬技能 > 软技能, 技能分类 + 细分数量
          - 成就质量 (20): 量化指标 + 影响范围 + 管理层级
          - 简历质量 (15): 完整度 + 结构化程度 + 猎头友好度
        
        预期分布: 60-80 中等, 81-90 优秀, 91-100 顶尖
        """
        import re as _re
        score = 0.0
        
        # ═══ 1. 教育背景 (max 15) ═══
        edu = structured.get("education", [])
        if edu:
            max_degree = 0
            for e in edu:
                deg = str(e.get("degree", "")).lower()
                if "博士" in deg or "phd" in deg or "doctor" in deg:
                    max_degree = max(max_degree, 10)
                elif "硕士" in deg or "master" in deg or "研究生" in deg:
                    max_degree = max(max_degree, 8)
                elif "本科" in deg or "bachelor" in deg:
                    max_degree = max(max_degree, 6)
                elif "大专" in deg or "专科" in deg:
                    max_degree = max(max_degree, 3)
            score += max_degree
            
            for e in edu:
                if e.get("is_qs50"):
                    score += 5
                    break
                elif e.get("is_985"):
                    score += 4
                    break
                elif e.get("is_211"):
                    score += 3
                    break
        else:
            score += 2
        
        # ═══ 2. 经验匹配 (max 25) ═══
        work = structured.get("work_experience", [])
        years = structured.get("years_experience", 0) or 0
        
        if 3 <= years <= 5:
            score += 6
        elif 6 <= years <= 10:
            score += 8
        elif 11 <= years <= 15:
            score += 6
        elif years > 15:
            score += 4
        elif years >= 1:
            score += 3
        
        if work:
            tiers = [w.get("company_tier", "") for w in work if w.get("company_tier")]
            if "FAANG" in tiers:
                score += 5
            elif "大厂" in tiers:
                score += 4
            elif "独角兽" in tiers:
                score += 3
            elif "上市公司" in tiers:
                score += 2
            
            team_sizes = [w.get("team_size", 0) for w in work if w.get("team_size")]
            if any(ts >= 20 for ts in team_sizes):
                score += 4
            elif any(ts >= 10 for ts in team_sizes):
                score += 3
            elif any(ts >= 5 for ts in team_sizes):
                score += 2
            
            stability = structured.get("career_stability", "")
            if stability == "稳定":
                score += 3
            elif stability == "一般":
                score += 1
        else:
            score += 2
        
        level = structured.get("role_level", "")
        if level:
            senior_kw = ["资深", "高级", "总监", "VP", "负责人", "主管", "经理", "head", "leader"]
            if any(kw in level.lower() for kw in senior_kw):
                score += 3
        
        # ═══ 3. 技能深度 (max 25) ═══
        skills = structured.get("skills", [])
        classified = structured.get("skills_classified", {})
        
        if skills:
            skill_count = len(skills)
            score += min(8, skill_count)
            
            domains = [k for k, v in classified.items() if isinstance(v, list) and v]
            domain_count = len(domains)
            score += min(6, domain_count * 2)
            
            ai_skills = classified.get("ai_ml", []) if isinstance(classified, dict) else []
            if ai_skills:
                ai_count = len(ai_skills)
                score += min(6, ai_count * 1.5)
            
            depth = structured.get("tech_depth", "")
            if depth == "深":
                score += 5
            elif depth == "中":
                score += 3
        else:
            score += 2
        
        # ═══ 4. 成就质量 (max 20) ═══
        quantified = 0
        if raw_text:
            quantified = len(_re.findall(
                r"\d+%|\d+人|\d+万|\d+亿|\d+x|[0-9]+\.[0-9]+x|降低\d+|提升\d+|增长\d+", raw_text))
        
        if quantified >= 8:
            score += 8
        elif quantified >= 5:
            score += 6
        elif quantified >= 3:
            score += 4
        elif quantified >= 1:
            score += 2
        
        projects = structured.get("projects", [])
        if projects:
            score += min(6, len(projects) * 2)
            if any(p.get("tech_stack") for p in projects):
                score += 3
            if any(p.get("impact") for p in projects):
                score += 3
        
        # ═══ 5. 简历质量 (max 15) ═══
        completeness = 0
        if structured.get("summary"): completeness += 3
        if structured.get("email") or structured.get("phone"): completeness += 2
        if skills: completeness += 2
        if work: completeness += 2
        if projects: completeness += 2
        score += completeness
        
        if raw_text:
            bullets = len(_re.findall(r"^[\s]*[-•·●◆▪►→]\s", raw_text, _re.MULTILINE))
            if bullets >= 10:
                score += 4
            elif bullets >= 5:
                score += 2
        
        return min(max(round(score, 1), 0), 100.0)

    def _rule_extract(self, text: str) -> dict:
        """Rule-based fallback extraction when LLM is unavailable."""
        result = {
            "candidate_name": "", "email": "", "phone": "", "wechat": "",
            "location": "", "current_role": "", "current_company": "",
            "years_experience": 0, "summary": "",
            "skills": [], "education": [], "work_experience": [],
            "projects": [], "certifications": [], "highlights": [],
            "skills_classified": {"ai_ml": [], "backend": [], "frontend": [],
                                   "cloud_devops": [], "data": [], "mobile": [],
                                   "product": [], "management": [], "other": []},
            "career_stability": "", "tech_depth": "", "industry_tags": [],
            "role_level": ""
        }
        if not text:
            return result

        lines = text.split('\n')
        text_lower = text.lower()

      # Name: extract Chinese name from ANY format (supports "I'M HULU 胡璐" / "John 张三")
        skip_kw = ['\u7b80\u5386', 'resume', '\u6c42\u804c', 'email', '@', 'http', 'tel', 'phone',
                   '\u6280\u80fd', '\u4e13\u4e1a\u6280\u80fd', '\u6838\u5fc3\u4f18\u52bf', '\u5de5\u4f5c\u7ecf\u9a8c', '\u9879\u76ee\u7ecf\u9a8c', '\u6559\u80b2\u80cc\u666f',
                   '\u81ea\u6211\u8bc4\u4ef7', '\u5de5\u4f5c\u7ecf\u5386', '\u4e2a\u4eba\u7b80\u4ecb', '\u610f\u5411', '\u6c42\u804c\u610f\u5411', '\u57fa\u672c\u4fe1\u606f',
                   '\u671f\u671b\u85aa\u8d44', '\u5230\u5c97\u65f6\u95f4', '\u4e2a\u4eba\u4f5c\u54c1', '\u5173\u4e8e\u6211', '\u7b80\u4ecb']
        for line in lines[:15]:
            line = line.strip()
            if not line:
                continue
            if any(kw in line for kw in skip_kw):
                continue
            # \u4f18\u5148\u4ece\u6df7\u5408\u683c\u5f0f\u4e2d\u63d0\u53d6\u4e2d\u6587\u540d (\u5982 "I\'M HULU \u80e1\u7490" \u2192 "\u80e1\u7490")
            cn = re.findall(r'[\u4e00-\u9fa5]{2,4}', line)
            if cn:
                result['candidate_name'] = cn[0]
                break
            # \u7eaf\u82f1\u6587\u540d\u56de\u9000 (only if no Chinese found in this line)
            en_match = re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?$', line.strip())
            if en_match:
                result['candidate_name'] = line.strip()
                break

        # Fallback: try extracting from lines with 2-4 pure Chinese chars
        if not result['candidate_name']:
            # Extended noise keywords that should NEVER be a name
            noise_names = {
                '专业技能', '教育经历', '工作经验', '项目经验', '工作经历',
                '个人简介', '自我评价', '自我介绍', '项目经历', '实习经历',
                '基本信息', '求职意向', '核心优势', '职业技能', '个人优势',
                '项目介绍', '主要工作', '负责内容', '关于我', '个人总结',
                '联系方式', '自我总结', '工作技能', '项目成果', '培训经历',
                '能力', '经验', '风险决策', '自我', '电话', '电话邮箱',
                '专业技能', '项目经历', '至小时', '至小', '至小时',
            }
            for line in lines[:20]:
                clean = re.sub(r'[^\u4e00-\u9fa5]', '', line)
                if 2 <= len(clean) <= 4:
                    if clean not in noise_names:
                        result['candidate_name'] = clean
                        break
            # If still not found, try the first line that is NOT noise
            if not result['candidate_name']:
                for line in lines[:20]:
                    clean = re.sub(r'[^\u4e00-\u9fa5]', '', line)
                    if 2 <= len(clean) <= 4 and len(clean) == len(line.strip()):
                        result['candidate_name'] = clean
                        break
        
        # Fallback 2: remove known noise prefixes from extracted name
        if result['candidate_name']:
            noise_prefixes = ['名：', '姓名：', '名字：', 'name：', 'Name：', '姓名:', '名字:']
            for prefix in noise_prefixes:
                if result['candidate_name'].startswith(prefix):
                    result['candidate_name'] = result['candidate_name'][len(prefix):]
                    break
            # Extract only Chinese chars if mixed
            cn = re.findall(r'[\u4e00-\u9fa5]', result['candidate_name'])
            if 2 <= len(cn) <= 4:
                result['candidate_name'] = ''.join(cn)

        # Email
        m = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', text)
        if m:
            result['email'] = m.group(0)

        # Phone (Chinese mobile + various formats)
        cleaned = text.replace(' ', '').replace('-', '').replace(' ', '').replace('\u200b', '')
        m = re.search(r'1[3-9]\d{9}', cleaned)
        if m:
            result['phone'] = m.group(0)
        else:
            # Try patterns like "手机：1xxxxxxxxxx" or "电话: 1xxxx" with spaces
            m = re.search(r'(?:手机|电话|mobile|tel|phone|Ｍｏｂｉｌｅ)[\s：: ]*1[\s3-9]', text, re.IGNORECASE)
            if m:
                # Extract following digits
                idx = text.find(m.group(0))
                surrounding = text[idx:idx+30]
                digits = re.sub(r'[^\d]', '', surrounding)
                if len(digits) >= 11:
                    result['phone'] = digits[:11]

        # WeChat
        m = re.search(r'(?:微信|wechat)[：:]\s*(\S+)', text, re.IGNORECASE)
        if m:
            result['wechat'] = m.group(1).strip()

        # Location
        locations = ['北京', '上海', '深圳', '杭州', '广州', '成都', '南京',
                     '武汉', '西安', '苏州', '天津', '重庆', '长沙']
        for loc in locations:
            if loc in text[:500]:
                result['location'] = loc
                break

        # Skills from known keywords
        skill_keywords = [
            'python', 'java', 'go', 'golang', 'rust', 'c++', 'c#', 'javascript',
            'typescript', 'react', 'vue', 'angular', 'node.js', 'nodejs',
            'django', 'flask', 'fastapi', 'spring', 'springboot',
            'kubernetes', 'k8s', 'docker', 'aws', 'gcp', 'azure',
            'mysql', 'postgresql', 'mongodb', 'redis', 'kafka', 'elasticsearch',
            'machine learning', 'deep learning', 'llm', 'rag', 'nlp',
            'tensorflow', 'pytorch', 'bert', 'gpt', 'langchain',
            'vector database', 'chromadb', 'product management',
            '数据分析', '项目管理', '敏捷开发', 'scrum',
        ]
        text_clean = text_lower.replace(',', ' ').replace('、', ' ').replace('，', ' ')
        found = []
        for skill in skill_keywords:
            if skill in text_clean:
                found.append(skill.title() if skill.isalpha() else skill)
        if found:
            result['skills'] = list(set(found))

        # Years of experience (multiple patterns)
        exp_pats = [
            r'(\d+)\s*年(?:经验|工作|开发|相关)',
            r'(?:经验|工作|开发)[：:\s]*(\d+)\s*年',
            r'(\d+)\s*years?\s*(?:of\s+|)(?:experience|exp)',
        ]
        for pat in exp_pats:
            m = re.search(pat, text_lower)
            if m:
                try:
                    result['years_experience'] = int(m.group(1))
                except ValueError:
                    pass
                break

        # Current role (simple heuristic: look for common role keywords)
        role_kw = ['工程师', '专家', '经理', '总监', '负责人', 'leader', 'manager', 'director']
        for line in lines[:30]:
            for kw in role_kw:
                if kw in line:
                    result['current_role'] = line.strip()[:40]
                    break
            if result['current_role']:
                break

        # Highlights
        hk = ['核心', '主导', '负责', '优化', '从0到1', '架构', '技术负责人', 'lead']
        for line in lines:
            for kw in hk:
                if kw in line and line.strip():
                    result['highlights'].append(line.strip()[:80])
                    break
            if len(result['highlights']) >= 3:
                break

        # Generate summary
        if result['candidate_name']:
            parts = [result['candidate_name']]
            if result['current_role']:
                parts.append(result['current_role'])
            if result['years_experience']:
                parts.append(str(result['years_experience']) + '年经验')
            if result['skills']:
                parts.append('熟悉' + '/'.join(result['skills'][:5]))
            result['summary'] = ' '.join(parts) + '。'

        return result
